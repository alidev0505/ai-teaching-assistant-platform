from docx import Document
from docx.shared import Pt, RGBColor
import io
import os
import re
import tempfile
import sys
from datetime import datetime

# Safely handle Windows-only COM libraries for cross-platform (Linux Azure) compatibility
try:
    if sys.platform.startswith('win'):
        import pythoncom
        import win32com.client
        from docx2pdf import convert
        WINDOWS_COM_AVAILABLE = True
    else:
        pythoncom = None
        win32com.client = None
        convert = None
        WINDOWS_COM_AVAILABLE = False
except ImportError:
    pythoncom = None
    win32com.client = None
    convert = None
    WINDOWS_COM_AVAILABLE = False

class ExportService:

    # ---------- 1. DOCUMENT SETTINGS ----------
    @staticmethod
    def get_document_meta(doc_type):
        dt = doc_type.lower().strip() if doc_type else ""
        if 'quiz' in dt:
            return {"marks": "10", "duration": "10 Minutes", "header": "Q1. Choose the best answers for the following MCQs. [10]"}
        elif 'assignment' in dt:
            return {"marks": "10", "duration": "N/A", "header": "Q1. Answer the following questions in detail. [10]"}
        elif 'mid' in dt:
            return {"marks": "20", "duration": "1 Hour", "header": "Q1. Attempt all questions. Section A (Objective) & Section B (Subjective)."}
        elif 'final' in dt:
            return {"marks": "40", "duration": "2 Hours", "header": "Q1. Attempt all questions. Section A (Objective) & Section B (Subjective)."}
        else:
            return {"marks": "10", "duration": "10 Mins", "header": "Questions:"}

    @staticmethod
    def get_department_name(program):
        p = str(program).upper()
        if "AI" in p or "ARTIFICIAL" in p:
            return "DEPARTMENT OF ARTIFICIAL INTELLIGENCE"
        elif "SE" in p or "SOFTWARE" in p:
            return "DEPARTMENT OF SOFTWARE ENGINEERING"
        elif "DS" in p or "DATA" in p:
            return "DEPARTMENT OF DATA SCIENCE"
        return "DEPARTMENT OF COMPUTER SCIENCE"

    # ---------- 2. TEXT CLEANER (REGEX POWERED) ----------
    @staticmethod
    def clean_text(text, include_answers=True):
        """Cleans AI output and optionally removes answers using secure Regular Expressions."""
        if not text: return ""

        text = text[:50000]

        text = re.sub(r'\*\*|__', '', text) 
        text = re.sub(r'\(#[^)]*don[^)]*t[^)]*use[^)]*topic[^)]*\)', '', text, flags=re.IGNORECASE)

        lines = text.split('\n')
        filtered_lines = []
        
        for line in lines:
            l = line.strip()
            if not l: continue 
            
            # 2. Aggressively strip leading symbols safely
            l = re.sub(r'^[\#\*\-\s]+', '', l)

            # 3. Skip Garbage Lines
            if re.match(r'^[-=_]{3,}$', l): continue
            if re.match(r'^(Title|Subject|Topic|Task|Date)\s*:', l, re.IGNORECASE): continue

            # 4. Skip Conversational Filler
            lower = l.lower()
            if lower.startswith("here is") or lower.startswith("sure, here") or \
               lower.startswith("as your professor") or "i have crafted" in lower:
                continue

            if not include_answers:
                if re.search(r'^(correct\s*)?(answer|option|ans)\s*[:\-]', lower):
                    continue
            
            filtered_lines.append(l)
        
        return "\n".join(filtered_lines)

    # ---------- 3. TEXT REPLACER ----------
    @staticmethod
    def replace_text_in_paragraph(paragraph, replacements):
        if not paragraph.text: return

        # 1. Try Simple Run Replacement
        for run in paragraph.runs:
            for key, val in replacements.items():
                if key in run.text:
                    run.text = run.text.replace(key, str(val))
                    run.bold = True
                    run.font.name = 'Times New Roman'
        
        # 2. Fallback Replacement
        text = paragraph.text
        updated = False
        for key, val in replacements.items():
            if key in text:
                text = text.replace(key, str(val))
                updated = True
        
        if updated:
            paragraph.text = text
            paragraph.style.font.name = 'Times New Roman'
            paragraph.style.font.size = Pt(11)
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                for val in replacements.values():
                    if str(val) in run.text:
                        run.bold = True

    # ---------- 4. MAIN GENERATOR ----------
    @staticmethod
    def create_university_doc(content_text, course_name, course_code, teacher_name, doc_type, program_name="BS-CS", include_answers=True):
        
        template_path = "Template.docx"
        if not os.path.exists(template_path):
            doc = Document()
            doc.add_paragraph("Error: Template.docx not found.")
            stream = io.BytesIO()
            doc.save(stream)
            stream.seek(0)
            return stream

        def sanitize_meta(val):
            return re.sub(r'[^\w\s\-\.\(\)\,\/:]', '', str(val or '')).strip()

        clean_course = sanitize_meta(course_name)
        clean_code = sanitize_meta(course_code)
        clean_teacher = sanitize_meta(teacher_name)
        clean_type = sanitize_meta(doc_type)
        clean_program = sanitize_meta(program_name)

        doc = Document(template_path)
        meta = ExportService.get_document_meta(clean_type)
        
        replacements = {
            "{{DEPARTMENT}}": ExportService.get_department_name(clean_program),
            "{{INSTRUCTOR}}": clean_teacher,
            "{{SESSION}}": f"{clean_type.capitalize()} (Fall-{datetime.now().year})",
            "{{SUBJECT}}": clean_course,
            "{{CODE}}": clean_code,
            "{{MARKS}}": meta["marks"],
            "{{PROGRAM}}": clean_program,
            "{{SHIFT}}": "Morning",
            "{{SECTION}}": "A",
            "{{DATE}}": datetime.now().strftime("%d %B %Y"),
            "{{DURATION}}": meta["duration"]
        }

        # 1. Fill Headers
        for p in doc.paragraphs:
            ExportService.replace_text_in_paragraph(p, replacements)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        ExportService.replace_text_in_paragraph(p, replacements)

        # 2. Fill Content
        content_filled = False
        question_counter = 1 

        for i, p in enumerate(doc.paragraphs):
            if "{{CONTENT}}" in p.text:
                p.text = "" 
                content_filled = True
                p.paragraph_format.space_after = Pt(0)

                run = p.add_run(meta["header"]) 
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

                if content_text:
                    cleaned_content = ExportService.clean_text(content_text, include_answers)
                    
                    for line in cleaned_content.split("\n"):
                        line = line.strip()
                        if not line: continue
                        
                        # --- A. DETECT ANSWERS (Backup Check) ---
                        is_answer = re.search(r'^(correct\s*)?(answer|option|ans)\s*[:\-]', line, re.IGNORECASE)
                        if is_answer:
                            if include_answers:
                                ans_p = doc.add_paragraph()
                                ans_p.paragraph_format.space_before = Pt(0)
                                ans_p.paragraph_format.space_after = Pt(6)
                                run = ans_p.add_run(line)
                                run.bold = True
                                run.font.color.rgb = RGBColor(0, 0, 0)
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(11)
                            continue 

                        # --- B. DETECT QUESTIONS ---
                        match = re.match(r'^(?:Question|Q)?\s*(\d+)[\.\):]\s*(.*)', line, re.IGNORECASE)
                        if match:
                            q_text = match.group(2).strip()
                            
                            if question_counter > 1:
                                empty_p = doc.add_paragraph()
                                empty_p.paragraph_format.space_after = Pt(12) 

                            q_label = doc.add_paragraph()
                            q_label.paragraph_format.space_after = Pt(0)
                            run = q_label.add_run(f"Question {question_counter}:")
                            run.bold = True
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(11)

                            if q_text:
                                q_body = doc.add_paragraph()
                                q_body.paragraph_format.space_after = Pt(6)
                                run = q_body.add_run(q_text)
                                run.bold = True 
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(11)
                            
                            question_counter += 1
                        
                        # --- C. DETECT SECTIONS ---
                        elif line.startswith("SECTION"):
                            doc.add_paragraph()
                            new_p = doc.add_paragraph()
                            run = new_p.add_run(line)
                            run.bold = True
                            run.font.name = 'Times New Roman'
                        
                        # --- D. DETECT OPTIONS ---
                        else:
                            p_opt = doc.add_paragraph()
                            p_opt.paragraph_format.space_after = Pt(0)
                            run = p_opt.add_run(line)
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(11)
                break 

        if not content_filled:
            doc.add_paragraph(meta["header"])
            doc.add_paragraph(content_text if content_text else "No Content Generated")

        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)
        return stream

    @staticmethod
    def docx_stream_to_pdf(docx_stream):
        if not WINDOWS_COM_AVAILABLE or not convert:
            print("PDF conversion via Word is only supported on Windows environments.")
            return None

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx:
            tmp_docx.write(docx_stream.getvalue())
            tmp_docx_path = tmp_docx.name
        
        tmp_pdf_path = tmp_docx_path.replace(".docx", ".pdf")
        try:
            pythoncom.CoInitialize()
            convert(tmp_docx_path, tmp_pdf_path)
            with open(tmp_pdf_path, "rb") as f:
                pdf_data = f.read()
            return io.BytesIO(pdf_data)
        except Exception as e:
            print(f"PDF Conversion Error: {e}")
            return None
        finally:
            if os.path.exists(tmp_docx_path): os.remove(tmp_docx_path)
            if os.path.exists(tmp_pdf_path): os.remove(tmp_pdf_path)
            if pythoncom:
                pythoncom.CoUninitialize()